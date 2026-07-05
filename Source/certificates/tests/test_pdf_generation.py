import os
import shutil
import subprocess
import tempfile
from datetime import date
from io import BytesIO
from pathlib import Path
from unittest import mock

from django.contrib.auth import get_user_model
from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse

from docx import Document

from Source.certificates.models import (
    AuditLog,
    Certificate,
    CertificateNumberSequence,
    Course,
    Learner,
)
from Source.certificates.services import (
    convert_docx_to_pdf,
    generate_certificate_pdf,
    render_certificate_docx,
)

User = get_user_model()

# A template exercising every placeholder the renderer fills.
TEMPLATE_TEXT = (
    "{{ learner_name }} | {{ id_number }} | {{ course_name }} | "
    "{{ certificate_number }} | {{ issue_date }} | {{ expiry_date }} | "
    "{{ assessor_name }}"
)


def docx_bytes(text):
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


class PdfGenerationTests(TestCase):
    def setUp(self):
        # Isolate uploaded templates / generated PDFs in a temp MEDIA_ROOT.
        self.media_root = tempfile.mkdtemp(prefix="media_")
        override = override_settings(MEDIA_ROOT=self.media_root)
        override.enable()
        self.addCleanup(override.disable)
        self.addCleanup(shutil.rmtree, self.media_root, True)

        self.admin = User.objects.create_superuser(
            username="admin", email="admin@example.com", password="pass12345"
        )
        self.client.force_login(self.admin)

        self.learner = Learner.objects.create(
            full_name="Thabo Mokoena",
            id_number="9001015800087",
            id_type="sa_id",
        )
        self.course = Course.objects.create(
            name="First Aid Level 1",
            level="Level 1",
            credits=5,
            saqa_id="119567",
            accreditation_body="HWSETA",
            validity_months=36,
            is_active=True,
            template_file=ContentFile(docx_bytes(TEMPLATE_TEXT), name="template.docx"),
        )

    def _make_certificate(self, certificate_number="0000001"):
        return Certificate.objects.create(
            certificate_number=certificate_number,
            learner=self.learner,
            course=self.course,
            learner_name_snapshot=self.learner.full_name,
            course_name_snapshot=self.course.name,
            training_date=date(2025, 1, 10),
            issue_date=date(2025, 1, 15),
            expiry_date=date(2028, 1, 15),
            assessor_name="Jane Assessor",
            status="issued",
            issued_by=self.admin,
        )

    def _add_url(self):
        return reverse("admin:certificates_certificate_add")

    def _post_data(self, **overrides):
        data = {
            "learner": self.learner.pk,
            "course": self.course.pk,
            "training_date": "2025-01-10",
            "issue_date": "2025-01-15",
            "assessor_name": "Jane Assessor",
        }
        data.update(overrides)
        return data

    def test_render_certificate_docx_fills_placeholders(self):
        cert = self._make_certificate()

        docx_path = render_certificate_docx(cert)
        try:
            text = "\n".join(p.text for p in Document(str(docx_path)).paragraphs)
        finally:
            shutil.rmtree(docx_path.parent, ignore_errors=True)

        self.assertIn("Thabo Mokoena", text)
        self.assertIn("9001015800087", text)
        self.assertIn("First Aid Level 1", text)

    def test_convert_docx_to_pdf_produces_pdf_file(self):
        tmp_dir = Path(tempfile.mkdtemp())
        try:
            docx_path = tmp_dir / "sample.docx"
            docx_path.write_bytes(docx_bytes("Hello World"))

            pdf_path = convert_docx_to_pdf(docx_path)

            self.assertTrue(pdf_path.exists())
            with open(pdf_path, "rb") as fh:
                self.assertEqual(fh.read(4), b"%PDF")
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def test_generate_certificate_pdf_attaches_file(self):
        cert = self._make_certificate()

        generate_certificate_pdf(cert)

        cert.refresh_from_db()
        self.assertTrue(cert.pdf_file.name)
        self.assertTrue(os.path.exists(cert.pdf_file.path))
        with cert.pdf_file.open("rb") as fh:
            self.assertEqual(fh.read(4), b"%PDF")

    def test_pdf_generation_failure_rolls_back_certificate(self):
        failure = subprocess.CalledProcessError(1, ["/usr/bin/soffice"])
        with mock.patch(
            "certificates.services.subprocess.run", side_effect=failure
        ):
            with self.assertRaises(subprocess.CalledProcessError):
                self.client.post(self._add_url(), self._post_data())

        # Nothing must survive a failed issuance -- not the cert, not the audit
        # log, and crucially not the certificate-number increment.
        self.assertEqual(Certificate.objects.count(), 0)
        self.assertEqual(AuditLog.objects.count(), 0)
        self.assertFalse(
            CertificateNumberSequence.objects.filter(last_number__gt=0).exists()
        )

    def test_template_file_rejects_non_docx(self):
        course = Course(
            name="Bad Course",
            level="Level 1",
            credits=5,
            saqa_id="000000",
            accreditation_body="HWSETA",
            validity_months=36,
            template_file=SimpleUploadedFile("template.odt", b"not a docx"),
        )
        with self.assertRaises(ValidationError):
            course.full_clean()
